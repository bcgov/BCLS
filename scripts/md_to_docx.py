"""Minimal Markdown -> Word converter for the Look West workplan.

Handles: ATX headings, GFM tables, ordered/unordered lists (incl. nested),
fenced code blocks, horizontal rules, and inline **bold**/*italic*/`code`.
Not a general Markdown engine -- just enough for this document.
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x00, 0x33, 0x66)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
CODE_GRAY = RGBColor(0x33, 0x33, 0x33)

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[.+?\]\(.+?\))")
LINK_RE = re.compile(r"\[(.+?)\]\((.+?)\)")


def add_inline(paragraph, text):
    """Add text to a paragraph, honoring **bold**, *italic*, `code`, [link](url)."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = CODE_GRAY
        elif part.startswith("["):
            m = LINK_RE.match(part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = NAVY
                run.underline = True
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)


def split_table_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def is_table_sep(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def main(src, dst):
    with open(src, encoding="utf-8") as fh:
        lines = fh.read().splitlines()

    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(10.5)

    i = 0
    n = len(lines)
    in_code = False
    code_buf = []

    while i < n:
        line = lines[i]

        # fenced code block
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = CODE_GRAY
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # horizontal rule
        if re.match(r"^---+\s*$", line):
            doc.add_paragraph().add_run("_" * 60).font.color.rgb = GOLD
            i += 1
            continue

        # table block
        if line.strip().startswith("|") and i + 1 < n and is_table_sep(lines[i + 1]):
            header = split_table_row(line)
            rows = []
            j = i + 2
            while j < n and lines[j].strip().startswith("|"):
                rows.append(split_table_row(lines[j]))
                j += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for idx, htext in enumerate(header):
                cell = table.rows[0].cells[idx]
                cell.paragraphs[0].clear()
                add_inline(cell.paragraphs[0], htext)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for row in rows:
                cells = table.add_row().cells
                for idx in range(len(header)):
                    txt = row[idx] if idx < len(row) else ""
                    cells[idx].paragraphs[0].clear()
                    add_inline(cells[idx].paragraphs[0], txt)
            doc.add_paragraph()
            i = j
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(20)
                run.font.color.rgb = NAVY
            elif level == 2:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = NAVY
            elif level == 3:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = GOLD
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
            i += 1
            continue

        # blockquote
        if line.strip().startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, line.strip()[1:].strip())
            i += 1
            continue

        # ordered list
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            p = doc.add_paragraph(style="List Number")
            if indent >= 2:
                p.paragraph_format.left_indent = Pt(36)
            add_inline(p, m.group(3))
            i += 1
            continue

        # unordered list
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 2:
                p.paragraph_format.left_indent = Pt(36)
            add_inline(p, m.group(2))
            i += 1
            continue

        # blank line
        if not line.strip():
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    doc.save(dst)
    print(f"Saved {dst}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
